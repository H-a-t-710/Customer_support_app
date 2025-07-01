import os
import re
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from app.core.config import settings
from app.utils.preprocessor import clean_html

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    Utility for loading and processing document files.
    Supports PDF and DOCX formats, with special handling for insurance documents.
    """
    
    def __init__(self, documents_dir: str = None):
        """
        Initialize document loader.
        
        Args:
            documents_dir (str, optional): Path to documents directory
        """
        self.documents_dir = documents_dir or settings.DOCUMENTS_PATH
    
    def load_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Load a document from file.
        
        Args:
            file_path (str): Path to document file
            
        Returns:
            List[Dict[str, Any]]: List of document parts with metadata
        """
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return []
        
        # Get file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # Process based on file type
            if file_ext == '.pdf':
                logger.info(f"Loading PDF document: {file_path}")
                return self._process_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                logger.info(f"Loading DOCX document: {file_path}")
                return self._process_docx(file_path)
            elif file_ext == '.txt':
                logger.info(f"Loading text document: {file_path}")
                return self._process_text(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_ext}")
                return []
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            return []
    
    def load_all_documents(self) -> List[Dict[str, Any]]:
        """
        Load all documents from the documents directory.
        
        Returns:
            List[Dict[str, Any]]: List of all document parts
        """
        all_documents = []
        
        try:
            # Check if directory exists
            if not os.path.exists(self.documents_dir):
                logger.warning(f"Documents directory not found: {self.documents_dir}")
                return []
            
            # Get all files in directory
            files = []
            
            # Walk through directory
            for root, _, filenames in os.walk(self.documents_dir):
                for filename in filenames:
                    # Check file extension
                    file_ext = os.path.splitext(filename)[1].lower()
                    if file_ext in ['.pdf', '.docx', '.doc', '.txt']:
                        file_path = os.path.join(root, filename)
                        files.append(file_path)
            
            logger.info(f"Found {len(files)} documents in {self.documents_dir}")
            
            # Process each file
            for file_path in files:
                logger.info(f"Processing document: {file_path}")
                document_parts = self.load_document(file_path)
                all_documents.extend(document_parts)
            
            logger.info(f"Loaded {len(all_documents)} document parts")
            
        except Exception as e:
            logger.error(f"Error loading documents from {self.documents_dir}: {str(e)}")
        
        return all_documents
    
    def _process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a PDF document.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            List[Dict[str, Any]]: List of document parts with metadata
        """
        document_parts = []
        filename = os.path.basename(file_path)
        
        try:
            # Open the PDF
            with fitz.open(file_path) as pdf:
                # Check if it's an SBC document (Summary of Benefits and Coverage)
                is_sbc = self._check_if_sbc(pdf)
                
                # Extract text from each page
                for page_num, page in enumerate(pdf):
                    # Extract text
                    page_text = page.get_text("text")
                    
                    # Handle tables differently for SBC documents
                    if is_sbc:
                        # Try to extract tables
                        tables = self._extract_tables_from_page(page)
                        
                        # If tables were found, process specially
                        if tables:
                            for table_num, table in enumerate(tables):
                                # Create document part for the table
                                document_parts.append({
                                    "content": table,
                                    "metadata": {
                                        "source": filename,
                                        "page": page_num + 1,
                                        "type": "pdf",
                                        "content_type": "table",
                                        "table_num": table_num + 1
                                    }
                                })
                    
                    # Create document part for the page text
                    if page_text.strip():
                        document_parts.append({
                            "content": page_text,
                            "metadata": {
                                "source": filename,
                                "page": page_num + 1,
                                "type": "pdf",
                                "content_type": "text"
                            }
                        })
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
        
        return document_parts
    
    def _process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a DOCX document.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            List[Dict[str, Any]]: List of document parts with metadata
        """
        document_parts = []
        filename = os.path.basename(file_path)
        
        try:
            # Open the document
            doc = Document(file_path)
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Create document part for paragraphs
            if paragraphs:
                document_parts.append({
                    "content": "\n\n".join(paragraphs),
                    "metadata": {
                        "source": filename,
                        "type": "docx",
                        "content_type": "text"
                    }
                })
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = self._extract_table_from_docx(table)
                if table_data:
                    document_parts.append({
                        "content": table_data,
                        "metadata": {
                            "source": filename,
                            "type": "docx",
                            "content_type": "table",
                            "table_num": table_num + 1
                        }
                    })
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        
        return document_parts
    
    def _process_text(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a text document.
        
        Args:
            file_path (str): Path to text file
            
        Returns:
            List[Dict[str, Any]]: List of document parts with metadata
        """
        document_parts = []
        filename = os.path.basename(file_path)
        
        try:
            # Read the file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Create document part
            document_parts.append({
                "content": text,
                "metadata": {
                    "source": filename,
                    "type": "text",
                    "content_type": "text"
                }
            })
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
        
        return document_parts
    
    def _check_if_sbc(self, pdf) -> bool:
        """
        Check if document is an SBC (Summary of Benefits and Coverage).
        
        Args:
            pdf: PyMuPDF document
            
        Returns:
            bool: True if document is an SBC
        """
        try:
            # Check first few pages for SBC indicators
            for i in range(min(3, len(pdf))):
                text = pdf[i].get_text("text").lower()
                if ("summary of benefits and coverage" in text or 
                    "what this plan covers" in text or
                    "what you will pay for covered services" in text or
                    "common medical event" in text):
                    return True
        except:
            pass
        
        return False
    
    def _extract_tables_from_page(self, page) -> List[str]:
        """
        Extract tables from PDF page with special handling for insurance documents.
        
        Args:
            page: PyMuPDF page
            
        Returns:
            List[str]: Extracted table content
        """
        tables = []
        
        try:
            # Try to get tables using built-in table detection
            tables_dict = page.find_tables()
            
            if tables_dict and tables_dict.tables:
                for table in tables_dict.tables:
                    # Extract cells
                    rows = []
                    for row_cells in table.extract():
                        rows.append(" | ".join([str(cell) for cell in row_cells]))
                    
                    # Construct table text
                    table_text = "\n".join(rows)
                    if table_text.strip():
                        tables.append(table_text)
            
            # Try to detect tables using layout heuristics if built-in detection didn't work
            if not tables:
                # Get page text with layout information
                html_text = page.get_text("html")
                
                # Clean up HTML
                cleaned_html = clean_html(html_text)
                
                # Check for table patterns
                if "| " in cleaned_html or " | " in cleaned_html:
                    # Split into lines
                    lines = cleaned_html.split("\n")
                    
                    # Find potential table sections
                    table_lines = []
                    current_table = []
                    in_table = False
                    
                    for line in lines:
                        if "|" in line:
                            current_table.append(line)
                            in_table = True
                        elif in_table and line.strip() and not line.strip().startswith("#"):
                            # Still part of the table description
                            current_table.append(line)
                        elif in_table:
                            # End of table
                            if len(current_table) > 2:  # At least a few lines to be a table
                                table_lines.append("\n".join(current_table))
                            current_table = []
                            in_table = False
                    
                    # Add the last table if there is one
                    if current_table and len(current_table) > 2:
                        table_lines.append("\n".join(current_table))
                    
                    # Add tables to result
                    tables.extend(table_lines)
        
        except Exception as e:
            logger.warning(f"Error extracting tables from page: {str(e)}")
        
        return tables
    
    def _extract_table_from_docx(self, table) -> str:
        """
        Extract a table from a DOCX document.
        
        Args:
            table: docx table object
            
        Returns:
            str: Table content as formatted string
        """
        rows = []
        
        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                text = cell.text.strip().replace('\n', ' ')
                cell_texts.append(text)
            
            rows.append(" | ".join(cell_texts))
        
        return "\n".join(rows) 